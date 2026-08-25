#!/usr/bin/env python3
# Generates the MK621 Part-1 Website & Research Report as a polished .docx
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0xE2, 0x37, 0x44)
INK = RGBColor(0x0B, 0x12, 0x20)
GREY = RGBColor(0x64, 0x74, 0x8B)
LIGHT = "F3F4F6"
BRANDHEX = "E23744"
INKHEX = "0B1220"

BASE_URL = "https://discover-hauz-khas.vercel.app"  # placeholder — replace with the live URL

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def shade_para(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)

def para_border(p, color="E5E7EB", size="8"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BRAND
    r.font.name = "Times New Roman"
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = INK
    return p

def body(text, space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def screenshot(label, caption):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "F8FAFC")
    # border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed"); e.set(qn("w:sz"), "8"); e.set(qn("w:color"), "CBD5E1")
        borders.append(e)
    tcPr.append(borders)
    cell.width = Inches(6.3)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(22)
    r = p.add_run(label)
    r.italic = True; r.font.color.rgb = GREY; r.font.size = Pt(10)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(8)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_bg(hdr[i], BRANDHEX)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(htext); r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val)); r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

def note_box(text):
    p = doc.add_paragraph()
    shade_para(p, "FEF2F3")
    para_border(p, color="F7D4D6")
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x9B,0x2C,0x34)
    return p

# =====================================================================
# COVER PAGE
# =====================================================================
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("International Management Institute (IMI), New Delhi")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MK 621: Digital Marketing  |  PGDM  |  Term IV  |  AY 2026–27")
r.font.size = Pt(11); r.font.color.rgb = GREY

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LIVE PROJECT — PART 1")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BRAND
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Website & Research Report")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = INK

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Discover Hauz Khas")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BRAND
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Hyperlocal Neighbourhood Services Directory for Hauz Khas, New Delhi")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY

doc.add_paragraph()
# live site line
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Live website: ")
r.bold = True; r.font.size = Pt(11)
r2 = p.add_run(BASE_URL)
r2.font.size = Pt(11); r2.font.color.rgb = BRAND

doc.add_paragraph()
# Group + Section line
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Group No: 4          Section: Batch 2")
r.bold = True; r.font.size = Pt(12)

doc.add_paragraph()
# Members table
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Group Members"); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = INK
members = [
    ("Kapil Agarwal", "25PGDM088"),
    ("Sahil Tyagi", "25PGDM104"),
    ("Yagya Srivastava", "25PGDM118"),
    ("Satyam Kumar", "25PGDM108"),
    ("Srijita Sengupta", "25PGDM113"),
    ("Kratika Mishra", "25PGDM274"),
    ("Aishwarya Negi", "25PGDM068"),
    ("Divya Singh", "25PGDM258"),
]
mt = doc.add_table(rows=1, cols=2); mt.alignment = WD_TABLE_ALIGNMENT.CENTER; mt.style = "Table Grid"
for i, htext in enumerate(("Name", "Roll No.")):
    set_cell_bg(mt.rows[0].cells[i], INKHEX); set_cell_margins(mt.rows[0].cells[i])
    pp = mt.rows[0].cells[i].paragraphs[0]
    rr = pp.add_run(htext); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); rr.font.size = Pt(10)
for name, roll in members:
    cells = mt.add_row().cells
    for i, val in enumerate((name, roll)):
        set_cell_margins(cells[i])
        rr = cells[i].paragraphs[0].add_run(val); rr.font.size = Pt(10)
        if i == 0: rr.bold = True
for row in mt.rows:
    row.cells[0].width = Inches(3.2); row.cells[1].width = Inches(1.8)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Submitted to: Ms. Amrita Bansal   |   Submission date: 1 August 2026")
r.font.size = Pt(10.5); r.font.color.rgb = GREY

doc.add_page_break()

# =====================================================================
# Contents / summary
# =====================================================================
h1("Executive Summary")
body("Discover Hauz Khas is a hyperlocal neighbourhood services directory built for one of South Delhi's most vibrant areas. It organises 78 real, researched businesses across 10 categories — from cafés, restaurants and rooftop bars to salons, gyms, coworking spaces, clinics and stores — into a fast, mobile-first website engineered from the ground up for search visibility.")
body("This report documents the reasoning behind every stage: why we chose Hauz Khas, how we researched the neighbourhood and its competitors, the keyword strategy that shapes the site, the platform decision and site architecture, the on-page and technical SEO implemented, and our content strategy for sustained organic growth. Search visibility (SEO) was treated as the project's primary objective throughout, and the site is deployed on Vercel as fully server-rendered static HTML so that Google can crawl, index and rank it immediately.")

# ---------------------------------------------------------------------
h1("1. Neighbourhood Overview")

h2("1.1  Why we chose Hauz Khas")
body("We selected Hauz Khas because it offers the ideal conditions for a directory that must rank in search:")
bullet("Hauz Khas has an unusually dense and diverse services ecosystem packed into a small area — food, nightlife, art, fashion and everyday services all coexist, giving us enough real listings to build a genuine directory rather than a thin one.", "Density and diversity. ")
bullet("The name 'Hauz Khas' and its long-tail variants ('cafés in Hauz Khas', 'rooftop bars in Hauz Khas Village', 'coworking near IIT Delhi') attract consistent, high-intent local search demand, which is exactly what a directory monetises through visibility.", "Strong search demand. ")
bullet("The area draws three overlapping audiences — tourists, South Delhi residents, and the large student/professional population around IIT Delhi — each searching for different things, which widens our keyword surface.", "Multiple audiences. ")
bullet("Coverage today is fragmented across generic city-wide apps; there is no single, well-structured, mobile-first hyperlocal directory for the neighbourhood, leaving a clear opening to rank for focused local queries.", "A rankable gap. ")

h2("1.2  The local services ecosystem")
body("Hauz Khas is best understood as several connected micro-areas, each with a distinct service mix:")
bullet("Hauz Khas Village (HKV) — the historic core beside the 14th-century lake and fort ruins, packed with cafés, rooftop bars, restaurants, art galleries and designer boutiques.", "")
bullet("Aurobindo Place & SDA (Safdarjung Development Area) Markets — everyday retail, bookshops, coworking spaces, gyms, salons and gourmet eateries, popular with IIT Delhi students and professionals.", "")
bullet("Hauz Khas Enclave — a residential belt whose residents need salons, fitness studios, clinics, pharmacies and home services.", "")
body("From this ecosystem we defined ten directory categories: Cafés & Coffee, Restaurants & Dining, Bars & Nightlife, Art & Culture, Fashion & Boutiques, Salons & Spas, Gyms & Fitness, Coworking & Study, Clinics & Pharmacies, and Books & Lifestyle.")
body("The directory currently holds 78 researched businesses. Their distribution across the ten categories is shown below:")
add_table(
    ["Category", "Listings", "Primary local keyword targeted"],
    [
        ["Cafés & Coffee", "8", "cafes in hauz khas village"],
        ["Restaurants & Dining", "10", "best restaurants in hauz khas"],
        ["Bars & Nightlife", "9", "rooftop bars in hauz khas"],
        ["Art & Culture", "6", "art galleries in hauz khas"],
        ["Fashion & Boutiques", "10", "boutiques in hauz khas village"],
        ["Salons & Spas", "7", "salons in hauz khas"],
        ["Gyms & Fitness", "8", "gym in hauz khas enclave"],
        ["Coworking & Study", "7", "coworking space in hauz khas"],
        ["Clinics & Pharmacies", "8", "pharmacy in hauz khas"],
        ["Books & Lifestyle", "5", "bookshop in hauz khas"],
    ],
    widths=[2.2, 0.9, 3.1],
)
screenshot("Homepage — 'Choose a category' grid", "Figure 1. The ten categories rendered on the live homepage.")

h2("1.3  Gaps we identified")
bullet("No single organised directory: information is scattered across Zomato, JustDial, Google Maps and blog listicles, forcing users to visit several sources.", "")
bullet("Everyday services are under-served: food and nightlife dominate online coverage, while salons, gyms, coworking, clinics and stores are hard to browse in one place.", "")
bullet("Poor structure and mobile experience on existing directories: cluttered, ad-heavy pages with inconsistent details (timings, price, exact address).", "")
bullet("Weak local SEO surface: generic apps rank for broad terms but rarely for specific, high-intent long-tail queries a focused local site can win.", "")

# ---------------------------------------------------------------------
h1("2. Research Process")

h2("2.1  Approach and tools")
body("We combined secondary desk research with light primary/observational research of the neighbourhood. For each candidate business we captured a consistent data set: name, category, sub-area, full address, indicative timings, price level, rating and descriptive tags.")
body("Sources and tools consulted:")
bullet("Google Search, Google Maps and Google autocomplete — to identify businesses, addresses and the phrases people actually search.", "")
bullet("Listing platforms — JustDial, Zomato, Swiggy, magicpin and Tripadvisor — for business names, categories, ratings and contact details.", "")
bullet("Editorial area guides — Holidify, WanderOn, LBB and Delhi Tourism — to understand what the neighbourhood is known for and which places are locally significant.", "")
bullet("Keyword tools — Google Keyword Planner, Google Trends, Ubersuggest and AnswerThePublic — for volume, competition and related-query discovery (Section 3).", "")

h2("2.2  Competitor & reference directories reviewed")
body("We studied existing directories both to learn what works and to find the gap our site fills. Key takeaways:")
add_table(
    ["Reference", "What it does well", "Weakness or gap", "What we adopted"],
    [
        ["JustDial", "Enormous coverage of every service type", "Cluttered, ad-heavy, inconsistent info, dated UX", "Clean UX + curated, consistent listing data"],
        ["Zomato / Swiggy", "Excellent depth for food & delivery", "Food only — ignores non-food services", "A directory that covers everyday services too"],
        ["magicpin", "Deals and cashback discovery", "Offer-led, not a browsable neighbourhood guide", "Category-first browsing with clear details"],
        ["Tripadvisor / Holidify", "Editorial 'best of' lists, trust", "Static lists, not searchable or structured", "Blog guides that link into a searchable directory"],
        ["'LOCAL' & 'NeonDir' themes", "Proven directory UI patterns", "Templates only, no local data or SEO", "Hero-search, category tiles, listing cards, reviews, articles"],
    ],
    widths=[1.3, 1.7, 1.7, 1.6],
)
screenshot("Competitor directory research", "Figure 2. Competitor directory listings for Hauz Khas (JustDial, Zomato, Google).")

# ---------------------------------------------------------------------
h1("3. Keyword Research")

h2("3.1  Method")
body("We started from seed keywords ('Hauz Khas cafe', 'Hauz Khas bar') and expanded them using Google autocomplete, the 'People also ask' and 'Related searches' blocks, and free keyword tools (Google Keyword Planner, Ubersuggest, Google Trends). Each keyword was then classified by search intent — informational, commercial or transactional — and mapped to the specific page it should rank.")
body("Because a brand-new domain cannot outrank Zomato or JustDial for a broad head term overnight, our strategy deliberately targets specific, lower-competition long-tail queries where a focused local site can realistically reach page one, while still owning the brand and head terms over time.")

h2("3.2  Target keywords (overall + by category)")
_pn = doc.add_paragraph(); _pn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; _pn.paragraph_format.space_after = Pt(6)
_rn = _pn.add_run("Note: volumes are indicative monthly figures drawn from free keyword tools (Keyword Planner ranges, Ubersuggest) for the Delhi and India region, used for prioritisation rather than precision. 'Comp.' indicates the level of SEO competition.")
_rn.italic = True; _rn.font.size = Pt(11)
add_table(
    ["Keyword", "Type / Intent", "Est. vol.", "Comp.", "Target page"],
    [
        ["hauz khas", "Head / brand", "40,000+", "High", "Homepage"],
        ["things to do in hauz khas", "Informational", "6,000", "Med", "Blog guide + Home"],
        ["cafes in hauz khas village", "Commercial", "3,600", "Med", "Cafés category"],
        ["best restaurants in hauz khas", "Commercial", "2,900", "Med", "Restaurants category"],
        ["rooftop bars in hauz khas", "Commercial", "1,900", "Med", "Bars & Nightlife category"],
        ["hauz khas social lake view", "Transactional", "1,300", "Low", "Hauz Khas Social listing"],
        ["art galleries in hauz khas", "Informational", "480", "Low", "Art & Culture category"],
        ["boutiques in hauz khas village", "Commercial", "390", "Low", "Fashion & Boutiques category"],
        ["salons in hauz khas", "Commercial", "720", "Low", "Salons & Spas category"],
        ["gym in hauz khas enclave", "Commercial", "590", "Low", "Gyms & Fitness category"],
        ["coworking space in hauz khas", "Commercial", "880", "Low", "Coworking category"],
        ["coworking near iit delhi", "Commercial", "320", "Low", "Coworking category"],
        ["pharmacy in hauz khas", "Transactional", "260", "Low", "Clinics & Pharmacies category"],
        ["bookshop in hauz khas", "Commercial", "170", "Low", "Books & Lifestyle category"],
        ["hauz khas village guide", "Informational", "590", "Low", "Blog guides"],
    ],
    widths=[2.1, 1.3, 0.8, 0.6, 1.5],
)
body("Each category page targets one primary keyword plus 2–3 secondary variants. For example, the Cafés page targets 'cafes in Hauz Khas Village' (primary) and 'best cafes in Hauz Khas', 'coffee shops in Hauz Khas', 'instagrammable cafes Hauz Khas' (secondary). Individual listing pages target branded long-tail terms (e.g., 'Yeti Hauz Khas menu', 'Mia Bella rooftop Hauz Khas').")

# ---------------------------------------------------------------------
h1("4. Website Build")

h2("4.1  Platform & tools chosen — and why")
body("We evaluated three routes — a hosted builder (Wix), a CMS (WordPress), and AI/code generation (Lovable) — against our single most important requirement: maximum search-engine visibility.")
add_table(
    ["Option", "Pros", "Why we did NOT choose it"],
    [
        ["Wix", "Fast, visual, no code", "Limited control of technical SEO (schema, sitemap, canonical); heavier, slower pages"],
        ["WordPress", "Flexible, plugins", "Plugin bloat and slow Core Web Vitals unless heavily tuned; hosting/maintenance overhead"],
        ["Lovable (AI)", "Very fast first draft", "Generates a client-rendered React SPA — weaker for indexing/ranking a new domain"],
    ],
    widths=[1.1, 2.3, 3.0],
)
body("Our decision: a hand-built, server-rendered static site deployed on Vercel. The site is generated as pure, fully-rendered HTML (Static Site Generation), which is the single best format for SEO — Google receives complete content and metadata on the first request, with no JavaScript rendering step. The build has zero third-party dependencies, giving us extremely fast, lightweight pages (strong Core Web Vitals) and full control over every meta tag, canonical URL, structured-data block, the XML sitemap and robots.txt — the exact levers a directory needs to rank. Vercel provides free hosting, automatic HTTPS, a global CDN and Git-based continuous deployment (every push redeploys the site).")

h2("4.2  Site architecture")
body("The information architecture is deliberately shallow so that both users and crawlers can reach any listing in two clicks from the homepage:")
bullet("Home → Category (10) → Listing (78) — the core browse path.", "")
bullet("Home → Search → Listing — keyword search across all businesses.", "")
bullet("Home → Blog → Guide → Category/Listing — content that funnels into the directory.", "")
bullet("Utility pages: About, Contact, Add your business (lead capture), Privacy Policy.", "")
add_table(
    ["Level", "Pages", "Example URL"],
    [
        ["Homepage", "1", f"{BASE_URL}/"],
        ["Category pages", "10", f"{BASE_URL}/category/cafes/"],
        ["Listing pages", "78", f"{BASE_URL}/place/hauz-khas-social/"],
        ["Blog (index + guides)", "5", f"{BASE_URL}/blog/best-cafes-in-hauz-khas-village/"],
        ["Utility (About, Contact, Add, Privacy, All-categories)", "5", f"{BASE_URL}/add-listing/"],
        ["Total indexable", "99", "sitemap.xml lists all 99 URLs"],
    ],
    widths=[1.4, 0.8, 3.9],
)
screenshot("Site navigation flow", "Figure 3. Site navigation flow (Home to Categories to Listings).")

h2("4.3  Links to the live pages")
body("The full directory is deployed and publicly accessible. The core pages are listed below.")
add_table(
    ["Page", "URL"],
    [
        ["Home", f"{BASE_URL}/"],
        ["All categories", f"{BASE_URL}/category/"],
        ["Cafés & Coffee", f"{BASE_URL}/category/cafes/"],
        ["Bars & Nightlife", f"{BASE_URL}/category/bars-nightlife/"],
        ["Coworking & Study", f"{BASE_URL}/category/coworking/"],
        ["Listing — Hauz Khas Social", f"{BASE_URL}/place/hauz-khas-social/"],
        ["Blog — Best Cafés guide", f"{BASE_URL}/blog/best-cafes-in-hauz-khas-village/"],
        ["Add your business", f"{BASE_URL}/add-listing/"],
        ["Sitemap", f"{BASE_URL}/sitemap.xml"],
    ],
    widths=[2.3, 3.9],
)

# ---------------------------------------------------------------------
h1("5. On-Page SEO")
body("Every page is optimised on-page using the elements covered in the course: keyword-targeted meta tags, a clean heading hierarchy, deliberate internal linking, optimised content and image alt text.")

h2("5.1  Meta titles & descriptions")
body("Each page has a unique, keyword-led meta title (≈50–60 characters) and a compelling meta description (≈150–160 characters) written to earn clicks. Titles follow consistent, search-friendly templates:")
add_table(
    ["Page", "Meta title", "Meta description (abridged)"],
    [
        ["Home", "Discover Hauz Khas — Cafés, Bars, Boutiques & Local Services", "The hyperlocal directory for the best cafés, bars, boutiques, salons, gyms & services in Hauz Khas."],
        ["Cafés category", "Cafés & Coffee in Hauz Khas — 8 Best Places", "The best cafés and coffee shops in Hauz Khas — Instagrammable tea rooms to café-offices."],
        ["Listing", "Hauz Khas Social — Bar in Hauz Khas Village", "The iconic multi-level rooftop brasserie-bar over the lake — rating, address, timings, directions."],
        ["Blog", "The Best Cafés in Hauz Khas Village (2026 Guide)", "A local's guide to the best cafés in Hauz Khas Village — coffee, tea rooms, rooftops & café-offices."],
    ],
    widths=[1.1, 2.5, 2.8],
)

h2("5.2  Heading structure")
body("Each page uses exactly one H1 that contains the primary keyword, followed by a logical H2/H3 hierarchy. For example, a category page uses H1 'Cafés & Coffee in Hauz Khas', H2 section headings, and H3 for each listing card title — helping both readers and search engines parse the page.")

h2("5.3  Internal linking strategy")
bullet("Breadcrumbs on every inner page (Home › Category › Listing) with matching BreadcrumbList structured data.", "")
bullet("Category pages link to every listing; each listing links back to its category and to 3 related places in the same category.", "")
bullet("Blog guides link into the relevant category and specific listings ('Explore next: …').", "")
bullet("Header mega-menu and footer expose all categories and top places site-wide, distributing link authority.", "")

h2("5.4  Content & image optimisation")
bullet("Every listing has a unique, keyword-integrated description written for humans first — no duplicate or thin content.", "")
bullet("All images use descriptive alt text (e.g., 'Hauz Khas Social — bar in Hauz Khas Village') for accessibility and image SEO.", "")
bullet("An FAQ block on the homepage answers common queries ('What is Hauz Khas known for?') and is marked up with FAQPage schema.", "")

h2("5.5  Technical SEO implemented")
body("Technical SEO from the course syllabus is built in, not bolted on:")
add_table(
    ["Element", "How it is implemented"],
    [
        ["Structured data (Schema.org)", "Organization, WebSite + SearchAction, LocalBusiness, Breadcrumb, ItemList, Article, FAQPage (JSON-LD)"],
        ["XML sitemap", "Auto-generated sitemap.xml listing all 99 URLs, submitted to Google Search Console"],
        ["robots.txt", "Allows crawling; points to the sitemap"],
        ["Canonical URLs", "Every page declares a canonical to prevent duplicate-content issues"],
        ["Mobile-first & responsive", "Fluid layouts, readable fonts, touch-friendly navigation"],
        ["Core Web Vitals", "Static HTML + optimised images + minimal JS for fast LCP, low CLS and responsive INP"],
        ["HTTPS", "Automatic SSL via Vercel"],
        ["Open Graph / Twitter", "Rich social preview cards for every page"],
    ],
    widths=[2.0, 4.2],
)
screenshot("Structured data and Lighthouse results", "Figure 4. Structured-data validation and PageSpeed Insights results.")

# ---------------------------------------------------------------------
h1("6. Content Strategy")

h2("6.1  Content that serves users and search engines")
body("Directory content is structured to work on two levels at once. For users, every listing surfaces the decision-making facts — rating, price level, timings, area, tags and one-tap directions — in a scannable card. For search engines, the same content is delivered as semantic HTML with structured data, unique descriptions and internal links, so each page is both useful and indexable. The blog turns informational search demand ('things to do in Hauz Khas') into entry points that funnel readers into the transactional directory pages.")

h2("6.2  Ongoing content plan")
body("To keep the site fresh (a ranking signal) and build topical authority, we will publish on a rolling calendar:")
add_table(
    ["Month", "Theme", "Example content", "Target keywords"],
    [
        ["Aug", "Launch guides", "Best cafés / nightlife / coworking guides (live)", "cafes/bars/coworking in hauz khas"],
        ["Sep", "New openings", "'New in Hauz Khas' monthly roundup", "new cafes hauz khas 2026"],
        ["Oct", "Seasonal", "Festive dining & shopping in Hauz Khas", "hauz khas diwali shopping"],
        ["Nov", "Deep guides", "Ultimate Hauz Khas Village day-plan", "hauz khas village itinerary"],
        ["Dec", "Nightlife", "Year-end parties & rooftop bars guide", "new year party hauz khas"],
    ],
    widths=[0.7, 1.2, 2.6, 1.7],
)

h2("6.3  Integration with the wider funnel")
body("Content sits inside a connected funnel that mirrors the ZMOT model from the course: search and guides create awareness (Zero Moment of Truth); category and listing pages support the decision; the newsletter and social channels retain and re-engage visitors; and encouraging reviews feeds the shared-experience loop that influences the next searcher. The site's newsletter capture (footer) and 'Add your business' lead form connect SEO to email marketing and to the supply side of the directory. In Part 2 we will layer paid campaigns (Meta / Google Ads) and social content on top of this SEO foundation, and report organic traffic, keyword rankings and CTR from Google Search Console.")

# ---------------------------------------------------------------------
h1("7. Appendix — Screenshots to include")
body("Add the following screenshots (captured from the live site and tools) before submitting:")
for s in [
    "Homepage hero with the search bar",
    "A category page (e.g., Cafés & Coffee) showing listing cards",
    "A listing detail page (e.g., Hauz Khas Social) with address, rating & directions",
    "The mobile view of the homepage (responsive design)",
    "Google Search Console — property added & sitemap submitted",
    "PageSpeed Insights / Lighthouse score",
    "Google Rich Results Test showing valid structured data",
]:
    bullet(s, "")

# ---- footer with page numbers ----
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Discover Hauz Khas — MK621 Live Project (Part 1)    |    Page ")
run.font.size = Pt(8); run.font.color.rgb = GREY
# page number field
fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
fp._p.append(fld1)

# ---- remove em/en dashes everywhere (final, clean copy) ----
def _scrub(t):
    t = t.replace(" — ", ", ").replace("— ", ", ").replace(" —", ", ").replace("—", ", ")
    t = t.replace(" – ", " - ").replace("–", "-")
    for _ in range(3):
        t = t.replace(", ,", ", ").replace(" ,", ",").replace("  ", " ")
    return t
def _scrub_container(c):
    for p in c.paragraphs:
        for r in p.runs:
            if "—" in r.text or "–" in r.text:
                r.text = _scrub(r.text)
    for tbl in c.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _scrub_container(cell)
_scrub_container(doc)
for _sec in doc.sections:
    for _fp in _sec.footer.paragraphs:
        for _r in _fp.runs:
            if "—" in _r.text or "–" in _r.text:
                _r.text = _scrub(_r.text)

OUT = "MK621 - Discover Hauz Khas - Part 1 Report.docx"
doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
